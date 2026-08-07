"""Document service for Word files.

The service organizes content extraction, asset publishing, and Markdown
assembly by file type on top of ``BaseDocumentService``.
"""

from __future__ import annotations

import base64
import logging
import zipfile
from pathlib import Path
from typing import Optional, TYPE_CHECKING

try:
    from docx import Document
except ImportError:
    Document = None

from .asset_reference import AssetReferenceMode, prepare_markdown_asset_references
from .base_document_service import ASSET_STAGE_NAMES, BaseDocumentService
from .document_artifact_models import DocumentAsset, MarkdownArtifact
from .document_artifact_writer import DocumentArtifactWriter
from .document_asset_publisher import (
    AftsDocumentAssetPublisher,
    DocumentAssetPublisher,
    NoOpDocumentAssetPublisher,
)
from .document_parse_logging import DocumentParseLogger
from .markdown_assembler import MarkdownAssembler, PlaceholderMarkdownAssembler

if TYPE_CHECKING:
    from services.afts_service import AftsService

from .paths import DOCUMENT_PARSE_WORKSPACE

logger = logging.getLogger(__name__)

class WordDocumentService(BaseDocumentService):
    """Provide document parsing for DOC and DOCX files."""

    _stage_names = ASSET_STAGE_NAMES
    _default_suffix = "docx"
    _empty_error_message = "Word 解析结果为空"

    def __init__(
        self,
        *,
        asset_reference_mode: AssetReferenceMode = "remote_id",
        markdown_assembler: Optional[MarkdownAssembler] = None,
        artifact_writer: Optional[DocumentArtifactWriter] = None,
    ) -> None:
        super().__init__(artifact_writer=artifact_writer)
        self._asset_reference_mode = asset_reference_mode
        self._markdown_assembler = markdown_assembler or PlaceholderMarkdownAssembler()

    async def _build_artifact(
        self,
        *,
        file_path: Path,
        task_id: str,
        source_file_name: str,
        afts_service: Optional["AftsService"],
        stage_logger: DocumentParseLogger,
    ) -> MarkdownArtifact:
        if Document is None:
            raise RuntimeError("未安装python-docx。请安装: pip install python-docx")

        output_dir = DOCUMENT_PARSE_WORKSPACE / task_id
        output_dir.mkdir(parents=True, exist_ok=True)

        with stage_logger.stage("content_extract"):
            markdown_text, assets = await self._extract_docx_content(
                file_path=file_path,
                output_dir=output_dir,
                source_file_name=source_file_name,
            )
            document_metrics = self._inspect_document(file_path)

        with stage_logger.stage(
            "asset_extract",
            asset_count=len(assets),
            afts_enabled=bool(afts_service),
        ):
            publisher = self._build_asset_publisher(afts_service)
            published_assets = await publisher.publish_assets(assets)
            prepare_markdown_asset_references(
                published_assets,
                output_dir=output_dir,
                asset_reference_mode=self._asset_reference_mode,
            )
            self._validate_published_assets(published_assets)

        artifact = MarkdownArtifact(
            markdown_text=markdown_text,
            assets=published_assets,
            diagnostics={
                "task_id": task_id,
                "source_file_name": source_file_name,
                "provider": "python_docx",
                "asset_count": len(published_assets),
                **document_metrics,
            },
        )
        with stage_logger.stage("markdown_assemble", asset_count=len(published_assets)):
            artifact.markdown_text = self._markdown_assembler.assemble(artifact)
        return artifact

    @staticmethod
    def _inspect_document(file_path: Path) -> dict[str, int]:
        document = Document(str(file_path))
        heading_count = sum(
            1
            for paragraph in document.paragraphs
            if str(getattr(getattr(paragraph, "style", None), "name", "")).lower().startswith("heading")
        )
        sections = list(document.sections)
        return {
            "paragraph_count": len(document.paragraphs),
            "heading_count": heading_count,
            "table_count": len(document.tables),
            "header_count": sum(1 for section in sections if section.header.paragraphs),
            "footer_count": sum(1 for section in sections if section.footer.paragraphs),
            "comment_count": len(getattr(document, "comments", []) or []),
        }

    def _build_asset_publisher(
        self,
        afts_service: Optional["AftsService"],
    ) -> DocumentAssetPublisher:
        if afts_service is None:
            return NoOpDocumentAssetPublisher()
        return AftsDocumentAssetPublisher(afts_service)

    async def _extract_docx_content(
        self,
        *,
        file_path: Path,
        output_dir: Path,
        source_file_name: str,
    ) -> tuple[str, list[DocumentAsset]]:
        doc = Document(str(file_path))
        images_dir = output_dir / f"{source_file_name}_images"
        assets = await self._extract_image_assets_from_docx(
            file_path=file_path,
            output_dir=images_dir,
        )
        content_parts: list[str] = []
        image_idx = 0

        for paragraph in doc.paragraphs:
            has_image = False
            for run in paragraph.runs:
                if run._element.xpath(".//a:blip"):
                    if image_idx < len(assets):
                        content_parts.append(self._build_asset_placeholder(assets[image_idx].asset_id))
                        content_parts.append("\n\n")
                    image_idx += 1
                    has_image = True

            paragraph_text = paragraph.text.strip()
            if paragraph_text:
                content_parts.append(self._convert_paragraph_to_markdown(paragraph_text, paragraph.style.name if paragraph.style else "Normal"))
            elif not has_image:
                continue

        if image_idx < len(assets):
            content_parts.append("\n## 图片\n\n")
            for asset in assets[image_idx:]:
                content_parts.append(self._build_asset_placeholder(asset.asset_id))
                content_parts.append("\n\n")

        if doc.tables:
            content_parts.append("\n## 表格\n\n")
            for table_idx, table in enumerate(doc.tables, start=1):
                content_parts.append(f"### 表格 {table_idx}\n\n")
                if not table.rows:
                    continue
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                content_parts.append("| " + " | ".join(headers) + " |\n")
                content_parts.append("|" + "---|" * len(headers) + "\n")
                for row in table.rows[1:]:
                    row_data = [cell.text.strip() for cell in row.cells]
                    content_parts.append("| " + " | ".join(row_data) + " |\n")
                content_parts.append("\n")

        return "".join(content_parts), assets

    async def _extract_image_assets_from_docx(
        self,
        *,
        file_path: Path,
        output_dir: Path,
    ) -> list[DocumentAsset]:
        assets: list[DocumentAsset] = []

        try:
            with zipfile.ZipFile(file_path, "r") as docx_zip:
                media_files = [name for name in docx_zip.namelist() if name.startswith("word/media/")]
                for order, media_file in enumerate(media_files, start=1):
                    original_name = Path(media_file).name
                    file_extension = Path(media_file).suffix.lower()
                    if file_extension not in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"}:
                        continue
                    try:
                        with docx_zip.open(media_file) as source:
                            image_data = source.read()
                    except BaseException as exc:
                        logger.warning(
                            "word_document_service failed to read media | media_file=%s error=%s",
                            media_file,
                            exc,
                            exc_info=True,
                        )
                        continue

                    asset = self._build_image_asset(
                        file_path=file_path,
                        output_dir=output_dir,
                        order=order,
                        original_name=original_name,
                        file_extension=file_extension,
                        image_data=image_data,
                    )
                    assets.append(asset)
        except BaseException as exc:
            logger.warning(
                "word_document_service failed to extract images from docx | file_path=%s error=%s",
                file_path,
                exc,
                exc_info=True,
            )
        return assets

    def _build_image_asset(
        self,
        *,
        file_path: Path,
        output_dir: Path,
        order: int,
        original_name: str,
        file_extension: str,
        image_data: bytes,
    ) -> DocumentAsset:
        mime_type = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".gif": "image/gif",
            ".bmp": "image/bmp",
            ".tiff": "image/tiff",
            ".webp": "image/webp",
        }.get(file_extension, "image/png")

        output_dir.mkdir(parents=True, exist_ok=True)
        local_path = output_dir / f"{file_path.stem}_img_{order - 1}{file_extension}"
        local_path.write_bytes(image_data)

        return DocumentAsset(
            asset_id=f"word_image_{order}",
            kind="embedded_image",
            remote_id="",
            local_path=local_path,
            order=order,
            meta={
                "index": str(order),
                "original_name": original_name,
                "mime_type": mime_type,
            },
        )

    def _validate_published_assets(
        self,
        assets: list[DocumentAsset],
    ) -> None:
        if self._asset_reference_mode == "local_path":
            return

        missing_assets = [
            asset for asset in assets if asset.local_path is not None and not asset.remote_id
        ]
        missing_remote_assets = [asset.asset_id for asset in missing_assets]
        if missing_remote_assets:
            local_asset_dirs = sorted(
                {
                    str(asset.local_path.parent)
                    for asset in missing_assets
                    if asset.local_path is not None
                }
            )
            local_asset_dir_message = ""
            if local_asset_dirs:
                local_asset_dir_message = (
                    "；本地图片目录: "
                    + ", ".join(local_asset_dirs)
                )
            raise RuntimeError(
                "DOCX 图片上传到 AFTS 失败，无法生成 markdown 图片引用: "
                + ", ".join(missing_remote_assets)
                + local_asset_dir_message
            )

    @staticmethod
    def _build_asset_placeholder(asset_id: str) -> str:
        return f"{{{{asset:{asset_id}}}}}"

    @staticmethod
    def _convert_paragraph_to_markdown(text: str, style_name: str) -> str:
        if "Heading 1" in style_name:
            return f"# {text}\n\n"
        if "Heading 2" in style_name:
            return f"## {text}\n\n"
        if "Heading 3" in style_name:
            return f"### {text}\n\n"
        if "Heading" in style_name:
            return f"#### {text}\n\n"
        return f"{text}\n\n"
